#!/usr/bin/env python3
"""Generate a bank-ready Word document from the business plan Markdown file."""

from __future__ import annotations

import argparse
import html
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "businessplan_robert_anna_walter.md"
OUTPUT = ROOT / "output/word/businessplan_robert_anna_walter_bankversion.docx"

INK = "24302B"
MUTED = "6F786E"
ACCENT = "657A52"
ACCENT_DARK = "3F5139"
LINE = "D8DDD4"
SOFT = "F3F5F0"
TABLE_ALT = "FAFBF8"
WHITE = "FFFFFF"

A4_WIDTH_CM = 21.0
A4_HEIGHT_CM = 29.7
CONTENT_WIDTH_DXA = 9504
TABLE_INDENT_DXA = 120


@dataclass
class TableBlock:
    rows: list[list[str]]
    aligns: list[str]


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def normalize_text(text: str) -> str:
    return (
        text.replace("\u00a0", " ")
        .replace("\u2011", "-")
        .replace("\ufeff", "")
        .strip()
    )


def clear_paragraph(paragraph) -> None:
    p_element = paragraph._p
    for child in list(p_element):
        p_element.remove(child)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 130, bottom: int = 90, end: int = 130) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, col_widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    tbl = table._tbl
    tbl_pr = tbl.tblPr

    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        existing = tbl_pr.find(qn(tag))
        if existing is not None:
            tbl_pr.remove(existing)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(col_widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)

    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in col_widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(0, grid)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, col_widths[index])


def set_table_borders(table, color: str = LINE, size: int = 4) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)
        borders.append(node)
    tbl_pr.append(borders)


def set_paragraph_border(paragraph, color: str = ACCENT, size: int = 12, space: int = 4) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)


def set_run_font(run, *, name: str = "Arial", size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_markdown_runs(paragraph, text: str, *, bold_default: bool = False, italic_default: bool = False, color: str | None = None, size: float | None = None) -> None:
    text = normalize_text(text)
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        content = part[2:-2] if is_bold else part
        run = paragraph.add_run(html.unescape(content))
        set_run_font(
            run,
            size=size,
            color=color,
            bold=bold_default or is_bold,
            italic=italic_default,
        )


def paragraph_spacing(paragraph, before: float = 0, after: float = 0, line: float | None = None) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if line is not None:
        fmt.line_spacing = line


def get_or_create_style(doc: Document, name: str, base: str | None = None):
    try:
        style = doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base:
        style.base_style = doc.styles[base]
    return style


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.22

    specs = {
        "Heading 1": (18.5, INK, 4, 12),
        "Heading 2": (13.5, ACCENT_DARK, 13, 6),
        "Heading 3": (11.2, INK, 8, 4),
        "Heading 4": (10.2, ACCENT_DARK, 6, 3),
    }
    for name, (size, color, before, after) in specs.items():
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    cover_label = get_or_create_style(doc, "Bank Cover Label", "Normal")
    cover_label.font.size = Pt(8.5)
    cover_label.font.bold = True
    cover_label.font.color.rgb = rgb(ACCENT_DARK)
    cover_label.paragraph_format.space_after = Pt(7)

    cover_title = get_or_create_style(doc, "Bank Cover Title", "Normal")
    cover_title.font.size = Pt(28)
    cover_title.font.bold = True
    cover_title.font.color.rgb = rgb(INK)
    cover_title.paragraph_format.space_after = Pt(9)
    cover_title.paragraph_format.line_spacing = 1.08

    cover_subtitle = get_or_create_style(doc, "Bank Cover Subtitle", "Normal")
    cover_subtitle.font.size = Pt(13.5)
    cover_subtitle.font.color.rgb = rgb(ACCENT_DARK)
    cover_subtitle.paragraph_format.space_after = Pt(19)
    cover_subtitle.paragraph_format.line_spacing = 1.2

    quote = get_or_create_style(doc, "Bank Quote", "Normal")
    quote.font.size = Pt(10.5)
    quote.font.italic = True
    quote.font.color.rgb = rgb(ACCENT_DARK)
    quote.paragraph_format.left_indent = Cm(0.35)
    quote.paragraph_format.right_indent = Cm(0.2)
    quote.paragraph_format.space_before = Pt(4)
    quote.paragraph_format.space_after = Pt(8)
    quote.paragraph_format.line_spacing = 1.2

    toc_title = get_or_create_style(doc, "Bank TOC Title", "Normal")
    toc_title.font.size = Pt(18)
    toc_title.font.bold = True
    toc_title.font.color.rgb = rgb(INK)
    toc_title.paragraph_format.space_after = Pt(10)

    meta = get_or_create_style(doc, "Bank Meta", "Normal")
    meta.font.size = Pt(9.4)
    meta.font.color.rgb = rgb(INK)
    meta.paragraph_format.space_after = Pt(0)
    meta.paragraph_format.line_spacing = 1.12

    table_cell = get_or_create_style(doc, "Bank Table Cell", "Normal")
    table_cell.font.size = Pt(8.6)
    table_cell.font.color.rgb = rgb(INK)
    table_cell.paragraph_format.space_after = Pt(0)
    table_cell.paragraph_format.line_spacing = 1.12

    table_header = get_or_create_style(doc, "Bank Table Header", "Bank Table Cell")
    table_header.font.bold = True
    table_header.font.color.rgb = rgb(WHITE)


def add_custom_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    next_abs = max(abstract_ids, default=20) + 1
    next_num = max(num_ids, default=20) + 1

    bullet_abs = next_abs
    decimal_abs = next_abs + 1
    bullet_num = next_num
    decimal_num = next_num + 1

    bullet_xml = f"""
    <w:abstractNum {nsdecls('w')} w:abstractNumId="{bullet_abs}">
      <w:multiLevelType w:val="singleLevel"/>
      <w:lvl w:ilvl="0">
        <w:start w:val="1"/>
        <w:numFmt w:val="bullet"/>
        <w:lvlText w:val="•"/>
        <w:lvlJc w:val="left"/>
        <w:pPr>
          <w:tabs><w:tab w:val="num" w:pos="540"/></w:tabs>
          <w:ind w:left="540" w:hanging="280"/>
          <w:spacing w:after="80" w:line="290" w:lineRule="auto"/>
        </w:pPr>
        <w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:hint="default"/><w:color w:val="{ACCENT_DARK}"/></w:rPr>
      </w:lvl>
    </w:abstractNum>
    """
    decimal_xml = f"""
    <w:abstractNum {nsdecls('w')} w:abstractNumId="{decimal_abs}">
      <w:multiLevelType w:val="singleLevel"/>
      <w:lvl w:ilvl="0">
        <w:start w:val="1"/>
        <w:numFmt w:val="decimal"/>
        <w:lvlText w:val="%1."/>
        <w:lvlJc w:val="left"/>
        <w:pPr>
          <w:tabs><w:tab w:val="num" w:pos="540"/></w:tabs>
          <w:ind w:left="540" w:hanging="280"/>
          <w:spacing w:after="80" w:line="290" w:lineRule="auto"/>
        </w:pPr>
      </w:lvl>
    </w:abstractNum>
    """
    numbering.append(parse_xml(bullet_xml))
    numbering.append(parse_xml(decimal_xml))

    for num_id, abstract_id in ((bullet_num, bullet_abs), (decimal_num, decimal_abs)):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    return bullet_num, decimal_num


def apply_numbering(paragraph, num_id: int, level: int = 0) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Seite ")
    set_run_font(run, size=8.0, color=MUTED)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    field_run = paragraph.add_run()
    field_run._r.append(fld_begin)
    field_run._r.append(instr)
    field_run._r.append(fld_end)
    set_run_font(field_run, size=8.0, color=MUTED)


def setup_document() -> tuple[Document, int, int]:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(A4_WIDTH_CM)
    section.page_height = Cm(A4_HEIGHT_CM)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.05)
    section.footer_distance = Cm(0.9)
    section.different_first_page_header_footer = True

    configure_styles(doc)
    bullet_num, decimal_num = add_custom_numbering(doc)

    header = section.header
    p = header.paragraphs[0]
    p.text = "Businessplan Robert & Anna Walter"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_spacing(p, after=1)
    for run in p.runs:
        set_run_font(run, size=8.0, color=MUTED)
    set_paragraph_border(p, color=LINE, size=4, space=2)

    footer = section.footer
    add_page_number(footer.paragraphs[0])
    return doc, bullet_num, decimal_num


def extract_cover_fields(lines: list[str]) -> dict[str, str]:
    title = "Businessplan"
    subtitle = ""
    founders = "Robert Walter & Anna Walter"
    found_title = False

    for index, line in enumerate(lines):
        stripped = normalize_text(line)
        if stripped.startswith("# ") and not found_title:
            title = stripped[2:].strip()
            found_title = True
        elif stripped.startswith("## ") and not subtitle:
            subtitle = stripped[3:].strip()
        elif stripped.startswith("**Gründer:**"):
            founders = stripped.replace("**Gründer:**", "").strip()
        if index > 12:
            break

    return {"title": title, "subtitle": subtitle, "founders": founders}


def add_spacer(doc: Document, points: float) -> None:
    paragraph = doc.add_paragraph()
    paragraph_spacing(paragraph, before=points, after=0)


def add_meta_table(doc: Document, fields: dict[str, str]) -> None:
    rows = [
        ("Gründer:innen", fields["founders"]),
        ("Standort", "Fliseryds-Boda, Mönsterås, Schweden"),
        ("Dokument", "Bankversion des Businessplans"),
        ("Stand", "24. Mai 2026"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [2150, 5900])
    set_table_borders(table, color=LINE, size=4)

    for row_index, (label, value) in enumerate(rows):
        for col_index, text in enumerate((label, value)):
            cell = table.cell(row_index, col_index)
            set_cell_shading(cell, SOFT)
            set_cell_margins(cell, top=115, bottom=115, start=150, end=150)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            clear_paragraph(paragraph)
            paragraph.style = "Bank Meta"
            paragraph_spacing(paragraph, after=0, line=1.12)
            add_markdown_runs(paragraph, text, bold_default=col_index == 0)


def add_cover(doc: Document, fields: dict[str, str]) -> None:
    add_spacer(doc, 58)

    rule = doc.add_paragraph()
    paragraph_spacing(rule, after=13)
    set_paragraph_border(rule, color=ACCENT, size=18, space=2)

    label = doc.add_paragraph(style="Bank Cover Label")
    add_markdown_runs(label, "Businessplan zur Bankeinreichung", color=ACCENT_DARK, size=8.5, bold_default=True)

    title = doc.add_paragraph(style="Bank Cover Title")
    add_markdown_runs(title, fields["title"], color=INK, size=28, bold_default=True)

    subtitle = doc.add_paragraph(style="Bank Cover Subtitle")
    add_markdown_runs(subtitle, fields["subtitle"], color=ACCENT_DARK, size=13.5)

    add_meta_table(doc, fields)
    add_spacer(doc, 62)

    closing = doc.add_paragraph(style="Bank Quote")
    add_markdown_runs(
        closing,
        "Erwerb und schrittweise Entwicklung eines naturnahen Erlebnis-, Bildungs- und Begegnungshofes mit konservativer Finanzierungs- und Umsetzungslogik.",
        italic_default=True,
        color=ACCENT_DARK,
        size=10.8,
    )
    set_paragraph_border(closing, color=ACCENT, size=8, space=6)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    title = doc.add_paragraph(style="Bank TOC Title")
    add_markdown_runs(title, "Inhaltsverzeichnis", color=INK, size=18, bold_default=True)

    rule = doc.add_paragraph()
    paragraph_spacing(rule, after=10)
    set_paragraph_border(rule, color=ACCENT, size=12, space=2)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Inhaltsverzeichnis in Word aktualisieren"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)
    set_run_font(run, size=9.5, color=MUTED)
    doc.add_page_break()


def body_lines_without_cover_intro(lines: list[str]) -> list[str]:
    for index, line in enumerate(lines):
        if normalize_text(line).startswith("# Teil 1"):
            return lines[index:]
    return lines


def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def is_table_separator(row: list[str]) -> bool:
    return bool(row) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in row)


def table_aligns(separator: list[str], count: int) -> list[str]:
    aligns: list[str] = []
    for cell in separator[:count]:
        cell = cell.strip()
        if cell.startswith(":") and cell.endswith(":"):
            aligns.append("center")
        elif cell.endswith(":"):
            aligns.append("right")
        else:
            aligns.append("left")
    while len(aligns) < count:
        aligns.append("left")
    return aligns


def table_widths(col_count: int) -> list[int]:
    if col_count == 2:
        return [6500, 2884]
    if col_count == 3:
        return [2700, 2500, 4184]
    return [CONTENT_WIDTH_DXA // col_count] * col_count


def add_table(doc: Document, block: TableBlock) -> None:
    col_count = max(len(row) for row in block.rows)
    rows = [row + [""] * (col_count - len(row)) for row in block.rows]
    widths = table_widths(col_count)
    table = doc.add_table(rows=len(rows), cols=col_count)
    set_table_geometry(table, widths)
    set_table_borders(table, color=LINE, size=4)

    for row_index, row in enumerate(rows):
        for col_index, text in enumerate(row):
            cell = table.cell(row_index, col_index)
            set_cell_margins(cell, top=95, bottom=95, start=130, end=130)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index == 0:
                set_cell_shading(cell, ACCENT_DARK)
            elif row_index % 2 == 0:
                set_cell_shading(cell, TABLE_ALT)

            paragraph = cell.paragraphs[0]
            clear_paragraph(paragraph)
            paragraph.style = "Bank Table Header" if row_index == 0 else "Bank Table Cell"
            paragraph_spacing(paragraph, after=0, line=1.12)
            if row_index > 0 and (block.aligns[col_index] == "right" or (col_index > 0 and re.search(r"\d", text))):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif row_index > 0 and block.aligns[col_index] == "center":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_markdown_runs(paragraph, text, bold_default=row_index == 0, color=WHITE if row_index == 0 else INK, size=8.6)

    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    after = doc.add_paragraph()
    paragraph_spacing(after, after=6)


def add_list_items(doc: Document, items: list[str], num_id: int) -> None:
    for item in items:
        paragraph = doc.add_paragraph()
        apply_numbering(paragraph, num_id)
        paragraph_spacing(paragraph, after=3, line=1.18)
        add_markdown_runs(paragraph, item, size=9.7)


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_spacing(paragraph, after=7, line=1.22)
    add_markdown_runs(paragraph, text, size=10.2)


def add_quote(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Bank Quote")
    add_markdown_runs(paragraph, text, italic_default=True, color=ACCENT_DARK, size=10.5)
    set_paragraph_border(paragraph, color=ACCENT, size=7, space=5)


def add_heading(doc: Document, level: int, text: str, first_h1: bool) -> bool:
    if level == 1:
        if not first_h1:
            doc.add_page_break()
        rule = doc.add_paragraph()
        paragraph_spacing(rule, after=8)
        set_paragraph_border(rule, color=ACCENT, size=14, space=2)

    style = f"Heading {min(level, 4)}"
    paragraph = doc.add_paragraph(style=style)
    add_markdown_runs(paragraph, text, bold_default=True)
    return False if level == 1 else first_h1


def parse_body(doc: Document, lines: list[str], bullet_num: int, decimal_num: int) -> None:
    paragraph_lines: list[str] = []
    list_items: list[str] = []
    ordered_items: list[str] = []
    first_h1 = True

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            add_body_paragraph(doc, " ".join(paragraph_lines))
            paragraph_lines = []

    def flush_lists() -> None:
        nonlocal list_items, ordered_items
        if list_items:
            add_list_items(doc, list_items, bullet_num)
            list_items = []
        if ordered_items:
            add_list_items(doc, ordered_items, decimal_num)
            ordered_items = []

    def flush_all() -> None:
        flush_paragraph()
        flush_lists()

    index = 0
    while index < len(lines):
        raw_line = lines[index].rstrip()
        line = normalize_text(raw_line)

        if not line:
            flush_all()
            index += 1
            continue

        if line == "---":
            flush_all()
            add_spacer(doc, 4)
            index += 1
            continue

        if line.startswith("|") and "|" in line[1:]:
            flush_all()
            table_lines: list[str] = []
            while index < len(lines):
                candidate = normalize_text(lines[index].rstrip())
                if not (candidate.startswith("|") and "|" in candidate[1:]):
                    break
                table_lines.append(candidate)
                index += 1

            rows = [split_table_row(row) for row in table_lines]
            if len(rows) >= 2 and is_table_separator(rows[1]):
                aligns = table_aligns(rows[1], len(rows[0]))
                table_rows = [rows[0], *rows[2:]]
            else:
                aligns = ["left"] * len(rows[0])
                table_rows = rows
            add_table(doc, TableBlock(table_rows, aligns))
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading_match:
            flush_all()
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            first_h1 = add_heading(doc, level, text, first_h1)
            index += 1
            continue

        bullet_match = re.match(r"^-\s+(.+)$", line)
        if bullet_match:
            flush_paragraph()
            if ordered_items:
                flush_lists()
            list_items.append(bullet_match.group(1))
            index += 1
            continue

        ordered_match = re.match(r"^\d+\.\s+(.+)$", line)
        if ordered_match:
            flush_paragraph()
            if list_items:
                flush_lists()
            ordered_items.append(ordered_match.group(1))
            index += 1
            continue

        if line.startswith(">"):
            flush_all()
            add_quote(doc, line.lstrip(">").strip())
            index += 1
            continue

        flush_lists()
        paragraph_lines.append(line)
        index += 1

    flush_all()


def mark_fields_dirty(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def build_docx(source: Path, output: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    fields = extract_cover_fields(lines)
    body_lines = body_lines_without_cover_intro(lines)

    doc, bullet_num, decimal_num = setup_document()
    core = doc.core_properties
    core.title = "Businessplan Robert & Anna Walter"
    core.author = fields["founders"]
    core.subject = "Bankversion"
    core.comments = "Aus Markdown generierte Word-Bankversion"

    add_cover(doc, fields)
    add_toc(doc)
    parse_body(doc, body_lines, bullet_num, decimal_num)
    mark_fields_dirty(doc)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main(argv: Iterable[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=SOURCE)
    parser.add_argument("--output", type=Path, default=OUTPUT)
    args = parser.parse_args(argv)

    build_docx(args.source, args.output)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
